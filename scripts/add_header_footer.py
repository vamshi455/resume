#!/usr/bin/env python3
"""Post-process a pandoc-generated docx to add header, footer, and watermark.

Usage: python3 scripts/add_header_footer.py <input.docx> [output.docx]
If output is omitted, overwrites the input file.
"""
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def add_page_field(paragraph, field_name="PAGE"):
    r = paragraph.add_run()
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    r._element.append(parse_xml('<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w')))
    r._element.append(parse_xml('<w:instrText %s xml:space="preserve"> %s </w:instrText>' % (nsdecls('w'), field_name)))
    r._element.append(parse_xml('<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w')))
    r._element.append(parse_xml('<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w')))


def dot_sep(paragraph):
    r = paragraph.add_run("  \u2022  ")
    r.font.size = Pt(7)
    r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)


def process(input_path, output_path=None):
    if output_path is None:
        output_path = input_path

    doc = Document(input_path)

    for section in doc.sections:
        section.different_first_page_header_footer = False
        section.header_distance = Emu(365760)
        section.footer_distance = Emu(365760)

        # ── HEADER ──
        header = section.header
        header.is_linked_to_previous = False

        for p in list(header.paragraphs):
            header._element.remove(p._element)
        for t in list(header.tables):
            header._element.remove(t._element)

        tbl = header.add_table(rows=1, cols=2, width=Inches(7.5))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True

        # Blue bottom border only
        tblPr = tbl._element.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = parse_xml('<w:tblPr %s/>' % nsdecls('w'))
            tbl._element.insert(0, tblPr)
        for old in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(old)
        tblPr.append(parse_xml(
            '<w:tblBorders %s>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="0077B5"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tblBorders>' % nsdecls('w')
        ))

        # Full width
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is not None:
            tblPr.remove(tblW)
        tblPr.append(parse_xml('<w:tblW %s w:type="pct" w:w="5000"/>' % nsdecls('w')))

        # Cell padding
        for cell in [tbl.cell(0, 0), tbl.cell(0, 1)]:
            tc = cell._element
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = parse_xml('<w:tcPr %s/>' % nsdecls('w'))
                tc.insert(0, tcPr)
            tcPr.append(parse_xml(
                '<w:tcMar %s>'
                '<w:top w:w="0" w:type="dxa"/><w:bottom w:w="20" w:type="dxa"/>'
                '<w:left w:w="0" w:type="dxa"/><w:right w:w="0" w:type="dxa"/>'
                '</w:tcMar>' % nsdecls('w')
            ))

        # LEFT: Name • Phone • Email
        lp = tbl.cell(0, 0).paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.paragraph_format.space_after = Pt(1)
        lp.paragraph_format.space_before = Pt(0)

        r = lp.add_run("Vamshi Singam")
        r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        dot_sep(lp)
        r = lp.add_run("+1 346-666-1192")
        r.font.size = Pt(7.5); r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        dot_sep(lp)
        r = lp.add_run("vamshi455@gmail.com")
        r.font.size = Pt(7.5); r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # RIGHT: LinkedIn URL • GitHub URL (plain text, no hyperlinks)
        rp = tbl.cell(0, 1).paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_after = Pt(1)
        rp.paragraph_format.space_before = Pt(0)

        r = rp.add_run("linkedin.com/in/vamshi-singam-82963556")
        r.font.size = Pt(7.5); r.font.color.rgb = RGBColor(0x00, 0x77, 0xB5)
        dot_sep(rp)
        r = rp.add_run("github.com/vamshi455")
        r.font.size = Pt(7.5); r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # ── FOOTER ──
        footer = section.footer
        footer.is_linked_to_previous = False

        for p in list(footer.paragraphs):
            footer._element.remove(p._element)
        for t in list(footer.tables):
            footer._element.remove(t._element)

        fp = footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(3)
        fp.paragraph_format.space_after = Pt(0)

        fpPr = fp._element.get_or_add_pPr()
        fpPr.append(parse_xml(
            '<w:pBdr %s><w:top w:val="single" w:sz="4" w:space="1" w:color="0077B5"/></w:pBdr>'
            % nsdecls('w')
        ))

        r = fp.add_run("Page ")
        r.font.size = Pt(7); r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        add_page_field(fp, "PAGE")
        r = fp.add_run(" of ")
        r.font.size = Pt(7); r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        add_page_field(fp, "NUMPAGES")

        # ── WATERMARK ──
        wm_p = header.add_paragraph()
        wm_p.paragraph_format.space_after = Pt(0)
        wm_p.paragraph_format.space_before = Pt(0)
        try:
            wm_p._element.append(parse_xml(
                '<w:r %s><w:rPr><w:noProof/></w:rPr><w:pict>'
                '<v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800" '
                'path="m@7,l@8,m@5,21600l@6,21600e" '
                'xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office">'
                '<v:formulas><v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/>'
                '<v:f eqn="sum 21600 0 @1"/><v:f eqn="sum 0 0 @2"/><v:f eqn="sum 21600 0 @3"/>'
                '<v:f eqn="if @0 @3 0"/><v:f eqn="if @0 21600 @1"/><v:f eqn="if @0 0 @2"/>'
                '<v:f eqn="if @0 @4 21600"/><v:f eqn="mid @5 @6"/><v:f eqn="mid @8 @5"/>'
                '<v:f eqn="mid @7 @8"/><v:f eqn="mid @6 @7"/><v:f eqn="sum @6 0 @5"/>'
                '</v:formulas>'
                '<v:path textpathok="t" o:connecttype="custom" '
                'o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800" o:connectangles="270,180,90,0"/>'
                '<v:textpath on="t" fitshape="t"/>'
                '<v:handles><v:h position="#0,bottomRight" xrange="6629,14971"/></v:handles>'
                '<o:lock v:ext="edit" text="t" shapetype="t"/>'
                '</v:shapetype>'
                '<v:shape id="PowerPlusWaterMarkObject" o:spid="_x0000_s2049" type="#_x0000_t136" '
                'style="position:absolute;margin-left:0;margin-top:0;width:500pt;height:100pt;'
                'rotation:315;z-index:-251658752;mso-position-horizontal:center;'
                'mso-position-horizontal-relative:margin;mso-position-vertical:center;'
                'mso-position-vertical-relative:margin" '
                'o:allowincell="f" fillcolor="#F0F0F0" stroked="f" '
                'xmlns:v="urn:schemas-microsoft-com:vml" xmlns:o="urn:schemas-microsoft-com:office:office">'
                '<v:fill opacity=".12"/>'
                '<v:textpath style="font-family:&quot;Calibri&quot;;font-size:1pt" string="CONFIDENTIAL"/>'
                '</v:shape></w:pict></w:r>' % nsdecls('w')
            ))
        except Exception:
            pass

    doc.save(output_path)
    print(f"Post-processed: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 scripts/add_header_footer.py <input.docx> [output.docx]")
        sys.exit(1)
    inp = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else None
    process(inp, out)
